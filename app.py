# IPD Case Sheet Generator Web Interface for Dr. Doodley Pet Hospital

import datetime
import streamlit as st
from docx import Document
from io import BytesIO

def generate_ipd_case_sheet_docx(data):
    doc = Document()
    doc.add_heading('Dr. Doodley Pet Hospital - Bangalore', 0)

    doc.add_paragraph(f"Pet Name: {data['pet_name']}     Pet ID: {data['pet_id']}")
    doc.add_paragraph(f"Date: {data['date']}     Time of Treatment: {data['time']}")

    doc.add_heading("Current Treatment Details", level=2)
    table = doc.add_table(rows=1, cols=7)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '#'
    hdr_cells[1].text = 'Name'
    hdr_cells[2].text = 'Route'
    hdr_cells[3].text = 'ml'
    hdr_cells[4].text = 'Dose'
    hdr_cells[5].text = 'Remarks'
    hdr_cells[6].text = 'Billed'
    for i, item in enumerate(data['injectables'], 1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = item['name']
        row[2].text = item['route']
        row[3].text = item['ml']
        row[4].text = item['dose']
        row[5].text = item['remarks']
        row[6].text = item['billed']

    doc.add_heading("Oral Medications", level=2)
    table = doc.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '#'
    hdr_cells[1].text = 'Name'
    hdr_cells[2].text = 'Dose'
    hdr_cells[3].text = 'Remarks'
    hdr_cells[4].text = 'Billed'
    for i, item in enumerate(data['orals'], 1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = item['name']
        row[2].text = item['dose']
        row[3].text = item['remarks']
        row[4].text = item['billed']

    doc.add_heading("Food Details", level=2)
    doc.add_paragraph(f"{data['food_type']} - {data['food_qty']}")
    if data['food_type'] == "Other" and data['food_other']:
        doc.add_paragraph(f"Other Food Details: {data['food_other']}")

    doc.add_heading("Emergency / Remarks", level=2)
    doc.add_paragraph(data['remarks'])

    doc.add_paragraph(f"\nTemp: {data['temp']}     CRT: {data['crt']}     Spo2: {data['spo2']}     BP: {data['bp']}")
    doc.add_paragraph(f"Doctor: {data['doctor']}     Paravet: {data['paravet']}")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

st.title("IPD Treatment Sheet Generator - Dr. Doodley Pet Hospital")

injectable_options = sorted([
    "ADRENALINE", "AMIKACIN", "AMOXY+SULBACTUM", "AMOXYCLAV", "APOMORPHINE", "ATROPINE", "AVIL",
    "BOTROPASE", "BUPRENORPHINE", "CALCIUM GLUCONATE", "CARPROFEN", "CEFTRIAXONE", "CEFTIOFUR",
    "CLINDAMYCIN", "CONVENIA", "CPM", "CYCLOPHOSPHAMIDE", "DARBOPOIETIN", "D5", "D25", "DEXAMETHASONE",
    "DEXMEDETOMIDINE", "DIAZEPAM", "DICYCLOMINE", "DERIPHYLIN", "DNS", "DOXOPRAM", "DOXYCYCLINE",
    "ETAMSYLATE", "ERYTHROPOEITIN", "FILGASTRIM", "FLUNIXIN", "FPP", "FRUSEMIDE", "GENTAMICIN",
    "GLYCOPYRROLATE", "HEPTOMAC", "IMIDOCARB", "INSULIN", "IVER MECTIN", "KETAMINE", "LEVERTRACETEM",
    "LIGNOCAINE", "MANNITOL", "MAROPITANT", "MELOXICAM", "MEROPENEM", "MIDAZOLAM", "METRIS", "NAC",
    "NER VIGEN", "NS", "ONDENSETRON", "PANTAPRAZOLE", "PERINORM", "PREDNISOLONE", "PROPOFOL", "RANTAC",
    "RL", "SOLUMEDROL", "STEMETIL", "THIOSOL", "TRENAXEMIC ACID", "TRIBIVET", "TRAMADOL", "VETALOG",
    "VINCRISTICINE", "VITAMIC C", "VITAMIN K", "XYLAZINE"])

oral_medications = sorted([
    "CK-Reno Feline", "CK-Reno", "Rinonadyl", "Cranbact", "Seveclear", "Pimoben 5/10", "Pimoben 1.25/2.5", "Kardioli",
    "Omnipulse", "Prolivet small", "Prolivet large", "Unomarin", "Ursopet", "Tocerapet", "Ocoxin", "EFA Pet", "Oclapet",
    "Atopivet Oral Suspension", "Atopivet Spot On", "Atosporin 50", "Atosporin 100", "Atosporin Oral Solution",
    "Terbopet 500", "Terbopet 250", "Terbopet-M", "Viv Keto Forte", "Viv Silky Shampoo", "Viv Silky Spray",
    "Viv Keto", "Pirofur", "Prurigo", "Subavet", "Relaxzyme small", "Relaxzyme large", "Spraid", "Impromune",
    "Viusid", "Clearify", "Otisoothe", "L-Sametine", "Alzer", "Diamel", "Vet Thyro 0.2", "Vet Thyro 0.6", "Asbrip",
    "Firotail 227", "Firotail 57", "Cartail 100", "Cartail 50", "Curcupet", "Cartail Spray", "Grapitail",
    "Entero Chronic", "Haltorin", "Viusid Detox", "Synfosium", "Pancreasolve", "Enterosecure BD", "Carminal",
    "Black Soldier Protein Powder", "Calowries", "Obex", "Condrovet Puppies", "Condrovet Force HA Large",
    "Condrovet Force HA", "Folrex", "Kalsis", "Vi Fi Spray", "Vi Fi Forte Spot On", "Exsite Shampoo", "Emepet",
    "Vet Maro tab", "Veticoal", "Prokolin", "Cephavet", "Aceptor", "Envas", "Carodyl 25 mg", "Carodyl 75 mg",
    "Carodyl 100mg", "Metaflam 1mg", "Metaflam 5 mg", "Ketochlor", "Lisybin small", "Lisybin medium", "Lisybin Large",
    "ProviBoost Drops", "ProviBoost Syrup", "Kiwof Plus", "Kiwof Plus XL", "Kiwof Puppy Suspension", "Kiwof cat",
    "Strongbeat advance", "Vetecto >4-10kg", "Vetecto 10-20 kg", "Vetecto20-40 kg", "Bravecto >4-10 kg",
    "Bravecto10-20 kg", "Bravecto 20-40 kg", "Bravecto> 40 kg", "Simparica", "Nutricoat Advance small",
    "Nutricoat Advance large", "Bioclan oral susp", "Bioclan 150 mg", "bioclan 300", "bioclan 600",
    "Clindapet oral susp", "Clindapet 150mg", "Clindapet 300", "Clindapet 600", "Zedoz 100", "Zedoz oral susp",
    "Zedoz 200", "Zedox 300", "Metrogyl 200 mg", "metrogyl 400mg", "metrogyl oral susp", "Samepet",
    "Samepet forte", "Ferritas", "aRBCe pet", "Thrombofit", "Advaplat", "Galibor", "Pronefra", "Cefpet syp",
    "Cefpet", "Cefpet XL", "Cefpet CLV", "Toxomox syp", "Toxomox 250", "Toxomox 500", "Fiprofort Plus 4-10 kg",
    "Fiprofort Plus 10-20 kg", "Fiprofort Plus 20-40kg", "Fiprofort Plus >40kg", "Fiprofort Spray", "CaniKuf",
    "Easibreath", "Otican", "Meo uripet intense", "uripet", "Uripet intense", "Smoothie pet derma", "Afoderm Hemp",
    "Pawflex", "Petjoint", "Apoquel", "Dermichlor", "Mycosan", "Ophthocare", "Ophthocare cool", "Ophthocare mono",
    "Ophthocare KT", "Ophthocare PD", "Ophthocare GP", "Ophthocare XT", "Ophthocare HYLO", "Auriko", "Pomisol",
    "Easotic", "Clearify", "Epiotic", "Ambiflush", "Digyton", "Digyton plus", "Mobility Plus", "Himpyrin",
    "Reliflam", "Aluspray", "D magg spray", "Diarest Cool", "Guttypet", "Gutwell", "Vendisc", "Poochrix", "Althromb",
    "E6WASH", "E6 lotion", "Interban", "Interban Maxima", "Interban LC", "Placentrix", "silverex", "Metlicho cat",
    "Metlicho dog", "Coatex", "Gabapentin 100", "Gabapentin 300"])

with st.form("ipd_form"):
    pet_name = st.text_input("Pet Name")
    pet_id = st.text_input("Pet ID")
    date = st.date_input("Date")
    time = st.time_input("Time of Treatment")

    injectables = []
    st.subheader("Current Treatment - Injectables")
    for i in range(1, 6):
        with st.expander(f"Injectable {i}"):
            name = st.selectbox(f"Name {i}", options=[""] + injectable_options, key=f"inj_name_{i}")
            route = st.text_input(f"Route {i}", key=f"inj_route_{i}")
            ml = st.text_input(f"ml {i}", key=f"inj_ml_{i}")
            dose = st.text_input(f"Dose {i}", key=f"inj_dose_{i}")
            remarks = st.text_input(f"Remarks {i}", key=f"inj_remark_{i}")
            billed = st.text_input(f"Billed {i}", key=f"inj_billed_{i}")
            if name:
                injectables.append({"name": name, "route": route, "ml": ml, "dose": dose, "remarks": remarks, "billed": billed})

    orals = []
    st.subheader("Oral Medications")
    for i in range(1, 4):
        with st.expander(f"Oral Med {i}"):
            name = st.text_input(f"Oral Name {i}", key=f"oral_name_{i}")
            dose = st.text_input(f"Dose {i}", key=f"oral_dose_{i}")
            remarks = st.text_input(f"Remarks {i}", key=f"oral_remarks_{i}")
            billed = st.text_input(f"Billed {i}", key=f"oral_billed_{i}")
            if name:
                orals.append({"name": name, "dose": dose, "remarks": remarks, "billed": billed})

    st.subheader("Food Details")
    food_options = ["Vivaldis Recovery Diet", "Vivaldis Dog GI", "Vivaldis Cat GI", "Other"]
    selected_food = st.selectbox("Type of Food", options=food_options)
    food_other = ""
    if selected_food == "Other":
        food_other = st.text_input("Specify Other Food")
    food_qty = st.text_input("Quantity")

    st.subheader("Emergency / Remarks")
    remarks = st.text_area("Remarks")

    temp = st.text_input("Temp")
    crt = st.text_input("CRT")
    spo2 = st.text_input("Spo2")
    bp = st.text_input("BP")

    doctor = st.selectbox("Doctor", ["Dr. Revathi", "Dr. Suhan", "Dr. Jyothi", "Dr. Jyothsna"])
    paravet = st.selectbox("Paravet", ["YASHWANTH", "RAKESH", "ANIL", "PARTHA", "VINAY", "PRAMOD", "PRAJWAL", "DARSHAN", "MAHESH"])

    submitted = st.form_submit_button("Generate Treatment Sheet")

if submitted:
    docx_file = generate_ipd_case_sheet_docx({
        "pet_name": pet_name,
        "pet_id": pet_id,
        "date": date,
        "time": time,
        "injectables": injectables,
        "orals": orals,
        "food_type": selected_food,
        "food_other": food_other,
        "food_qty": food_qty,
        "remarks": remarks,
        "temp": temp,
        "crt": crt,
        "spo2": spo2,
        "bp": bp,
        "doctor": doctor,
        "paravet": paravet
    })

    st.success("Treatment sheet generated successfully!")

    st.download_button(
        label="Download Treatment Sheet (Word)",
        data=docx_file,
        file_name="Treatment_Sheet.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
